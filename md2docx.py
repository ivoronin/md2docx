import argparse
import io
import sys
from importlib import import_module

from docx import Document
from docx.enum.text import WD_BREAK
import mistune

class DocXMarkdown(mistune.Markdown):
    def parse(self, text):
        s = io.BytesIO()
        super().parse(text)
        self.renderer.doc.save(s)
        return s.getvalue()
        

class DocXRenderer(mistune.Renderer):
    def __init__(self, style, *largs, **kargs):
        self.doc = Document()
        if style:
            style.apply(self.doc)
        self.clist = None
        super().__init__(*largs, **kargs)

    def list_item(self, text):
        return f"{text}\n"
      
    def list(self, body, ordered=True):
        if ordered:
            style = "ListNumber3"
        else:
            style = "ListBullet"
        for i in body.rstrip().split("\n"):
            self.doc.add_paragraph(i, style)
        return ''
        
    def header(self, text, level, raw=None):
        self.doc.add_heading(text, level)
        return ''

    def hrule(self):
        self.doc.paragraphs[-1].runs[-1].add_break(WD_BREAK.PAGE)
        return ''

    def block_quote(self, text):
        self.doc.paragraphs[-1].style = "Quote"
        return ''

    def paragraph(self, text):
        self.doc.add_paragraph(text, style=None)
        return ''

def parse_args(args):
    parser = argparse.ArgumentParser(description="Converts a markdown document into docx")
    parser.add_argument("input", help="Markdown input file")
    parser.add_argument("output", help="docx output file (will overwrite if it already exists)")
    parser.add_argument('--style', '-s', default='default', help="Style name to use")
    args  = parser.parse_args()
    return args
                            

def main():
    args = parse_args(sys.argv[1:])
    if args.style:
        try:
            style = import_module(f'styles.{args.style}').Style
        except ModuleNotFoundError:
            print(f"{sys.argv[0]}: style {args.style} is not found")
            sys.exit(1)
    else:
        style = None
    renderer = DocXRenderer(style)
    markdown = DocXMarkdown(renderer = renderer)
    with open(args.input) as f:
        with open(args.output, "wb") as g:
            g.write(markdown(f.read()))


if __name__ == '__main__':
    main()
    
