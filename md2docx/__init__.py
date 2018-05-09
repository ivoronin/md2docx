"""
md2docs - renders markdown as Microsoft Word documents
"""
import argparse
import io
import sys
import re
from importlib import import_module
from docx import Document
from docx.enum.text import WD_BREAK
import mistune
from lxml import etree


class DocXWriter:
    """Parses html, writes docx"""
    def __init__(self, output, style=None):
        self._doc = Document()
        self._output = output
        if style:
            style.apply(self._doc)
        self._cur_para = None
        self._para_style_stack = [None]
        self._list_level = 0

    def parse(self, html):
        """Parses supplied html string"""
        parser = etree.XMLParser(remove_blank_text=True)
        tree = etree.XML('<html>' + html + '</html>', parser=parser)
        self._walk(tree)
        self._doc.save(self._output)

    @property
    def _next_para_style(self):
        assert self._para_style_stack
        return self._para_style_stack[-1]

    def _walk(self, root):
        for elem in root:
            # At element start
            if elem.tag in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                level = int(elem.tag[1])
                self._doc.add_heading(elem.text, level=level)
            elif elem.tag == 'blockquote':
                self._para_style_stack.append('Quote')
            elif elem.tag == 'p':
                self._cur_para = self._doc.add_paragraph(elem.text, style=self._next_para_style)
            elif elem.tag == 'hr':
                self._doc.paragraphs[-1].runs[-1].add_break(WD_BREAK.PAGE)
            elif elem.tag in ['ul', 'ol']:
                self._list_level += 1
                assert self._list_level in range(1, 4)
                style_suffix = '' if self._list_level == 1 else f' {self._list_level}'
                style_kind = 'Bullet' if elem.tag == 'ul' else 'Number'
                style_name = f'List {style_kind}{style_suffix}'
                self._para_style_stack.append(style_name)
            elif elem.tag == 'li':
                self._cur_para = self._doc.add_paragraph(elem.text, style=self._next_para_style)
            elif elem.tag == 'strong':
                self._cur_para.add_run(elem.text, style='Strong')
            elif elem.tag == 'em':
                self._cur_para.add_run(elem.text, style='Emphasis')
            else:
                raise RuntimeError(f"Unexpected tag {elem.tag}")

            # Recursively walk
            self._walk(elem)

            # At element end
            if elem.tag in ['ul', 'ol']:
                self._list_level -= 1
                self._para_style_stack.pop()
            elif elem.tag == 'blockquote':
                self._para_style_stack.pop()
            elif elem.tag in ['p', 'ul']:
                self._cur_para = None

            # Add all non-whitespace tails to current paragraph
            if elem.tail and not elem.tail.isspace():
                self._cur_para.add_run(elem.tail)


def walk(root):
    for element in root:
        print(element.tag, element.text, element.tail)
        walk(element)


def parse_args(args):
    """
    Parses command line arguments
    Args:
        args (list): Command line arguments
    Returns:
        Namespace: object holding attributes
    """
    parser = argparse.ArgumentParser(description="Converts a markdown document into docx")
    parser.add_argument('input', help="Markdown input file")
    parser.add_argument('output', help="docx output file (will overwrite if it already exists)")
    parser.add_argument('--style', '-s', default='default', help="Style name to use")
    args = parser.parse_args()
    return args


def main():
    """
    Main entry point
    """
    args = parse_args(sys.argv[1:])
    if args.style:
        try:
            style = import_module(f'{__name__}.styles.{args.style}').Style
        except ModuleNotFoundError:
            print(f"{sys.argv[0]}: style {args.style} is not found")
            sys.exit(1)
    else:
        style = None
    markdown = mistune.Markdown(renderer=mistune.Renderer(use_xhtml=1))
    with open(args.input) as source:
        html = markdown(source.read())
    DocXWriter(args.output, style).parse(html)

if __name__ == '__main__':
    main()
